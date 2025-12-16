"""
Enhanced Document Processor với OCR cải tiến (FIX tiếng Việt bị rụng chữ)
- Render PDF ở DPI cao + sharpen nhẹ
- Preprocessing (denoise + contrast + adaptive threshold) + deskew
- PaddleOCR gọi bằng ocr.ocr(...) ổn định cho scan
- Normalize Unicode NFC để tránh rơi dấu
- Export Word set font đầy đủ để không mất tiếng Việt
"""

import os
import re
import cv2
import fitz
import numpy as np
import unicodedata
from PIL import Image, ImageEnhance

from paddleocr import PaddleOCR
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from typing import List, Dict, Tuple, Any


# =========================
# Utils: Unicode
# =========================
def nfc(text: str) -> str:
    return unicodedata.normalize("NFC", text) if text else text


def vi_cleanup(text: str) -> str:
    """Cleanup nhẹ nhàng sau OCR cho tiếng Việt."""
    if not text:
        return text
    text = nfc(text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\s+([,.;:!?])", r"\1", text)
    text = re.sub(r"([,.;:!?])([^\s])", r"\1 \2", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# =========================
# Image Preprocessor
# =========================
class ImagePreprocessor:
    """Tiền xử lý hình ảnh trước khi OCR"""

    @staticmethod
    def enhance_image(image: np.ndarray) -> np.ndarray:
        """Cải thiện chất lượng hình ảnh cho OCR."""
        if image is None:
            return image

        # grayscale
        if len(image.shape) == 3:
            gray = cv2.cvtColor(image, cv2.COLOR_RGB2GRAY)
        else:
            gray = image

        # denoise
        denoised = cv2.fastNlMeansDenoising(gray, None, 12, 7, 21)

        # contrast (CLAHE)
        clahe = cv2.createCLAHE(clipLimit=2.2, tileGridSize=(8, 8))
        contrast = clahe.apply(denoised)

        # adaptive threshold
        binary = cv2.adaptiveThreshold(
            contrast, 255,
            cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY,
            31, 10
        )

        # morph close (nhẹ)
        kernel = np.ones((1, 1), np.uint8)
        cleaned = cv2.morphologyEx(binary, cv2.MORPH_CLOSE, kernel)

        return cleaned

    @staticmethod
    def deskew_image(image: np.ndarray) -> np.ndarray:
        """Xoay thẳng hình ảnh bị nghiêng."""
        try:
            if image is None:
                return image

            # cần binary để tìm góc
            img = image
            if len(img.shape) == 3:
                img = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)

            coords = np.column_stack(np.where(img < 255))  # chữ thường là điểm tối
            if coords.shape[0] < 50:
                return image

            angle = cv2.minAreaRect(coords)[-1]
            if angle < -45:
                angle = -(90 + angle)
            else:
                angle = -angle

            if abs(angle) < 0.3:
                return image

            (h, w) = image.shape[:2]
            center = (w // 2, h // 2)
            M = cv2.getRotationMatrix2D(center, angle, 1.0)
            rotated = cv2.warpAffine(
                image, M, (w, h),
                flags=cv2.INTER_CUBIC,
                borderMode=cv2.BORDER_REPLICATE
            )
            return rotated
        except Exception:
            return image

    @staticmethod
    def resize_for_ocr(image: np.ndarray, target_height: int = 2200) -> np.ndarray:
        """Resize image về kích thước tối ưu cho OCR."""
        if image is None:
            return image

        h, w = image.shape[:2]
        if h <= 0 or w <= 0:
            return image

        if h < target_height:
            scale = target_height / h
            new_w = int(w * scale)
            new_h = target_height
            return cv2.resize(image, (new_w, new_h), interpolation=cv2.INTER_CUBIC)

        if h > target_height * 1.6:
            scale = target_height / h
            new_w = int(w * scale)
            new_h = target_height
            return cv2.resize(image, (new_w, new_h), interpolation=cv2.INTER_AREA)

        return image


# =========================
# Vietnamese Spell Checker (fix position = char offset)
# =========================
class VietnameseSpellChecker:
    """Kiểm tra chính tả tiếng Việt (đơn giản)."""

    def __init__(self):
        self.dictionary = set([
            'và', 'của', 'có', 'trong', 'là', 'được', 'với', 'các', 'cho', 'để',
            'đã', 'này', 'theo', 'những', 'người', 'từ', 'một', 'năm', 'khi', 'về',
            'nội', 'dung', 'hình', 'ảnh', 'đồ', 'thị', 'tài', 'liệu', 'chương', 'mục',
            'phần', 'bảng', 'số', 'liệu', 'thông', 'tin', 'dữ', 'kiểm', 'tra', 'xử', 'lý',
            'phương', 'pháp', 'kết', 'quả', 'nghiên', 'cứu', 'phát', 'triển', 'hệ', 'thống',
        ])
        self.load_extended_dictionary()

    def load_extended_dictionary(self, dict_file='vietnamese_dict.txt'):
        if os.path.exists(dict_file):
            with open(dict_file, 'r', encoding='utf-8') as f:
                words = f.read().splitlines()
                self.dictionary.update([w.strip().lower() for w in words if w.strip()])

    def check_text(self, text: str) -> Dict[str, Any]:
        text = nfc(text or "")
        # dùng finditer để lấy luôn char offset
        pattern = re.compile(
            r'[\wàáạảãâầấậẩẫăằắặẳẵèéẹẻẽêềếệểễìíịỉĩòóọỏõôồốộổỗơờớợởỡùúụủũưừứựửữỳýỵỷỹđ]+',
            re.IGNORECASE
        )

        errors = []
        total = 0

        for m in pattern.finditer(text.lower()):
            total += 1
            word = m.group(0)
            if len(word) > 2 and word not in self.dictionary:
                errors.append({
                    'word': word,
                    'position': m.start(),     # ✅ char offset (để sửa theo vị trí)
                    'suggestions': []
                })

        return {
            'total_words': total,
            'errors': errors,
            'error_rate': (len(errors) / total) if total > 0 else 0
        }


# =========================
# Main Processor
# =========================
class EnhancedDocumentProcessor:
    """Document processor với OCR cải tiến"""

    def __init__(self, output_dir='output'):
        print("🔧 Đang khởi tạo PaddleOCR...")
        print("   (Lần đầu sẽ tải model, có thể mất vài phút...)")

        # ✅ dùng API ổn định: ocr.ocr(...)
        self.ocr = PaddleOCR(
            lang='vi',
            use_angle_cls=True,
            show_log=False
        )

        self.preprocessor = ImagePreprocessor()
        self.spell_checker = VietnameseSpellChecker()
        self.output_dir = output_dir

        os.makedirs(output_dir, exist_ok=True)
        os.makedirs(f"{output_dir}/images", exist_ok=True)
        os.makedirs(f"{output_dir}/debug", exist_ok=True)

        self.results = {
            'sections': [],
            'images': [],
            'spelling_check': {},
            'metadata': {}
        }

        print("✅ PaddleOCR sẵn sàng!\n")

    # ---------
    # OCR Page
    # ---------
    def process_page_image(self, image_rgb: np.ndarray, page_num: int, debug: bool = False) -> List[Dict[str, Any]]:
        """Xử lý một trang: preprocessing + OCR + sort theo layout."""
        print(f"   🔍 Đang OCR trang {page_num}...")

        if image_rgb is None:
            return []

        # 1) resize -> deskew -> enhance (binary)
        img = self.preprocessor.resize_for_ocr(image_rgb, target_height=2200)
        img = self.preprocessor.deskew_image(img)

        # PaddleOCR nhận ảnh RGB/BGR đều được, nhưng preproc binary là 1 kênh
        bin_img = self.preprocessor.enhance_image(img)

        if debug:
            debug_path = os.path.join(self.output_dir, "debug", f"page_{page_num:03d}_bin.png")
            cv2.imwrite(debug_path, bin_img)

        # 2) OCR
        try:
            ocr_result = self.ocr.ocr(bin_img, cls=True)  # ✅ ổn định nhất
        except Exception as e:
            print(f"   ❌ Lỗi OCR: {e}")
            return []

        if not ocr_result:
            print(f"   ⚠️  Không phát hiện text trong trang {page_num}")
            return []

        layout_elements: List[Dict[str, Any]] = []
        kept = 0

        # ocr_result: list[line] where line = [box, (text, score)]
        for line in ocr_result:
            try:
                box = line[0]
                text, confidence = line[1][0], float(line[1][1])

                if confidence < 0.50:
                    continue

                text = vi_cleanup(text)
                if not text:
                    continue

                # bbox
                x_coords = [p[0] for p in box]
                y_coords = [p[1] for p in box]
                x_min, x_max = float(min(x_coords)), float(max(x_coords))
                y_min, y_max = float(min(y_coords)), float(max(y_coords))
                width = x_max - x_min
                height = y_max - y_min

                element_type = self.classify_element(text, width, height, y_min)

                layout_elements.append({
                    'type': element_type,
                    'bbox': {
                        'x': int(x_min),
                        'y': int(y_min),
                        'width': int(width),
                        'height': int(height)
                    },
                    'text': text,
                    'confidence': confidence
                })
                kept += 1
            except Exception:
                continue

        # 3) sort top->bottom, left->right (để ghép text đúng thứ tự)
        layout_elements.sort(key=lambda e: (e['bbox']['y'], e['bbox']['x']))

        print(f"   ✅ Giữ {kept} dòng (confidence > 0.5)")
        return layout_elements

    def classify_element(self, text: str, width: float, height: float, y_pos: float) -> str:
        if height > 32 and y_pos < 220:
            return 'title'
        if re.match(r'^\d+\.\d+', text.strip()):
            return 'heading'
        return 'text'

    # -------------------
    # Extract Images PDF
    # -------------------
    def extract_images_from_pdf(self, pdf_path: str) -> List[Dict[str, Any]]:
        doc = fitz.open(pdf_path)
        images = []

        for page_num in range(len(doc)):
            page = doc[page_num]
            image_list = page.get_images()

            for img_idx, img in enumerate(image_list):
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                image_ext = base_image["ext"]

                img_filename = f"img_{page_num+1:03d}_{img_idx+1:03d}.{image_ext}"
                img_path = os.path.join(self.output_dir, "images", img_filename)

                with open(img_path, "wb") as f:
                    f.write(image_bytes)

                img_rects = page.get_image_rects(xref)
                bbox = {'x': 0, 'y': 0, 'width': 0, 'height': 0}

                if img_rects:
                    rect = img_rects[0]
                    bbox = {
                        'x': int(rect.x0),
                        'y': int(rect.y0),
                        'width': int(rect.width),
                        'height': int(rect.height)
                    }

                images.append({
                    'id': f'img_{page_num+1:03d}_{img_idx+1:03d}',
                    'filename': img_filename,
                    'path': img_path,
                    'page': page_num + 1,
                    'format': image_ext.upper(),
                    'bbox': bbox
                })

        doc.close()
        return images

    # -------------
    # Process PDF
    # -------------
    def process_pdf(self, pdf_path: str, debug: bool = False) -> Dict[str, Any]:
        print(f"\n{'='*80}")
        print(f"📄 ĐANG XỬ LÝ: {os.path.basename(pdf_path)}")
        print(f"{'='*80}\n")

        print("🖼️  BƯỚC 1: Trích xuất hình ảnh...")
        self.results['images'] = self.extract_images_from_pdf(pdf_path)
        print(f"✅ Đã trích xuất {len(self.results['images'])} hình ảnh\n")

        print("🔍 BƯỚC 2: OCR văn bản...")
        doc = fitz.open(pdf_path)

        all_text_parts: List[str] = []
        section_counter = 1
        total_pages = len(doc)

        for page_idx in range(total_pages):
            page_num = page_idx + 1
            print(f"\n📄 Trang {page_num}/{total_pages}")
            page = doc[page_idx]

            # Render page: tăng zoom + sharpen nhẹ bằng PIL
            zoom = 3.5
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat, alpha=False)

            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            # sharpen nhẹ
            img = ImageEnhance.Sharpness(img).enhance(1.5)
            img_array = np.array(img)

            layout_elements = self.process_page_image(img_array, page_num, debug=debug)

            current_section = ""

            for element in layout_elements:
                element['page'] = page_num

                if element['type'] == 'heading':
                    if current_section.strip():
                        self.results['sections'].append({
                            'id': f'section_{section_counter}',
                            'content': vi_cleanup(current_section) + '\n\n</break>\n',
                            'page': page_num
                        })
                        section_counter += 1
                        current_section = ""

                    current_section = f"\n{element['text']}\n\n"
                else:
                    current_section += element['text'] + " "

                all_text_parts.append(element['text'])

            if current_section.strip():
                self.results['sections'].append({
                    'id': f'section_{section_counter}',
                    'content': vi_cleanup(current_section),
                    'page': page_num
                })
                section_counter += 1

        doc.close()

        print(f"\n📝 BƯỚC 3: Kiểm tra chính tả...")
        all_text = vi_cleanup(" ".join(all_text_parts))
        self.results['spelling_check'] = self.spell_checker.check_text(all_text)

        self.results['metadata'] = {
            'total_pages': total_pages,
            'total_sections': len(self.results['sections']),
            'total_images': len(self.results['images']),
            'total_words': self.results['spelling_check']['total_words'],
            'spelling_errors': len(self.results['spelling_check']['errors'])
        }

        print(f"\n{'='*80}")
        print("✅ HOÀN TẤT XỬ LÝ")
        print(f"{'='*80}")
        print(f"📊 Tổng số trang: {self.results['metadata']['total_pages']}")
        print(f"📝 Tổng số phần: {self.results['metadata']['total_sections']}")
        print(f"🖼️  Tổng số hình: {self.results['metadata']['total_images']}")
        print(f"📖 Tổng số từ: {self.results['metadata']['total_words']}")

        return self.results

    # ----------------
    # Export to Word
    # ----------------
    def export_to_word(self, output_filename: str = 'output_document.docx') -> str:
        print(f"\n💾 Đang xuất file Word: {output_filename}")

        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(13)

        # ✅ set font đủ để Word không rớt tiếng Việt
        rFonts = style._element.rPr.rFonts
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        rFonts.set(qn('w:cs'), 'Times New Roman')

        title = doc.add_heading('TÀI LIỆU ĐÃ XỬ LÝ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_heading('Thông tin tổng quan', 1)
        md = self.results.get('metadata', {})
        metadata_text = (
            f"Tổng số trang: {md.get('total_pages','N/A')}\n"
            f"Tổng số phần: {md.get('total_sections','N/A')}\n"
            f"Tổng số hình ảnh: {md.get('total_images','N/A')}\n"
            f"Tổng số từ: {md.get('total_words','N/A')}\n"
            f"Lỗi chính tả: {md.get('spelling_errors','N/A')}\n"
        )
        doc.add_paragraph(metadata_text)

        doc.add_heading('Nội dung', 1)

        for section in self.results.get('sections', []):
            content = section.get('content', '')
            content = vi_cleanup(content)

            if re.match(r'^\d+\.\d+', content.strip()):
                lines = content.split('\n', 1)
                heading_text = lines[0].strip()
                doc.add_heading(heading_text, 2)

                if len(lines) > 1:
                    body_text = lines[1].replace('</break>', '').strip()
                    body_text = vi_cleanup(body_text)
                    if body_text:
                        p = doc.add_paragraph(body_text)
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            else:
                text = content.replace('</break>', '').strip()
                text = vi_cleanup(text)
                if text:
                    p = doc.add_paragraph(text)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            if '</break>' in content:
                doc.add_page_break()

        if self.results.get('images'):
            doc.add_page_break()
            doc.add_heading('Danh sách hình ảnh', 1)

            for img in self.results['images']:
                doc.add_heading(f"Hình {img['id']}", 2)
                try:
                    doc.add_picture(img['path'], width=Inches(4))
                except Exception:
                    doc.add_paragraph(f"[Không thể thêm hình ảnh: {img['filename']}]")

        output_path = os.path.join(self.output_dir, output_filename)
        doc.save(output_path)
        print(f"✅ Đã lưu file Word: {output_path}\n")
        return output_path


def main():
    pdf_path = "input_document.pdf"
    if not os.path.exists(pdf_path):
        print(f"❌ Không tìm thấy file: {pdf_path}")
        print("📝 Đặt file PDF cần xử lý vào thư mục hiện tại và đổi tên thành 'input_document.pdf'")
        return

    processor = EnhancedDocumentProcessor(output_dir='output')
    results = processor.process_pdf(pdf_path, debug=True)
    processor.export_to_word('tai_lieu_da_xu_ly.docx')


if __name__ == "__main__":
    main()
