"""
Công cụ Xử lý Tài liệu Nâng cao
- PaddleOCR cho OCR
- Layout detection với bounding boxes
- Kiểm tra lỗi chính tả tiếng Việt
- Trích xuất hình ảnh
- Xuất file Word với cấu trúc

Cài đặt thư viện:
pip install paddleocr paddlepaddle pdf2image pillow python-docx opencv-python pyspellchecker
pip install fitz PyMuPDF reportlab
"""

import os
import cv2
import fitz  # PyMuPDF
import numpy as np
from PIL import Image
from paddleocr import PaddleOCR
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import re
from spellchecker import SpellChecker
from typing import List, Dict, Tuple
import json


class VietnameseSpellChecker:
    """Kiểm tra chính tả tiếng Việt"""
    
    def __init__(self):
        # Từ điển tiếng Việt cơ bản (có thể mở rộng)
        self.dictionary = set([
            'và', 'của', 'có', 'trong', 'là', 'được', 'với', 'các', 'cho', 'để',
            'đã', 'này', 'theo', 'những', 'người', 'từ', 'một', 'năm', 'khi', 'về',
            'nội', 'dung', 'hình', 'ảnh', 'đồ', 'thị', 'tài', 'liệu', 'chương', 'mục',
            'phần', 'bảng', 'số', 'liệu', 'thông', 'tin', 'dữ', 'kiểm', 'tra', 'xử', 'lý',
            'phương', 'pháp', 'kết', 'quả', 'nghiên', 'cứu', 'phát', 'triển', 'hệ', 'thống',
            'ứng', 'dụng', 'công', 'nghệ', 'khoa', 'học', 'giáo', 'dục', 'kinh', 'tế',
            'xã', 'hội', 'văn', 'hóa', 'chính', 'trị', 'quốc', 'gia', 'dân', 'tộc',
            'cấp', 'độ', 'mức', 'tăng', 'giảm', 'cao', 'thấp', 'lớn', 'nhỏ', 'mới', 'cũ'
        ])
        
        # Load từ điển mở rộng từ file nếu có
        self.load_extended_dictionary()
    
    def load_extended_dictionary(self, dict_file='vietnamese_dict.txt'):
        """Load từ điển mở rộng từ file"""
        if os.path.exists(dict_file):
            with open(dict_file, 'r', encoding='utf-8') as f:
                words = f.read().splitlines()
                self.dictionary.update(words)
    
    def check_word(self, word: str) -> Tuple[bool, List[str]]:
        """Kiểm tra một từ và đưa ra gợi ý"""
        word_lower = word.lower()
        
        # Bỏ qua số và ký tự đặc biệt
        if not re.match(r'^[\wàáạảãâầấậẩẫăằắặẳẵèéẹẻẽêềếệểễìíịỉĩòóọỏõôồốộổỗơờớợởỡùúụủũưừứựửữỳýỵỷỹđ]+$', word_lower):
            return True, []
        
        # Từ quá ngắn
        if len(word_lower) <= 2:
            return True, []
        
        # Kiểm tra trong từ điển
        if word_lower in self.dictionary:
            return True, []
        
        # Tìm từ gợi ý
        suggestions = self.find_suggestions(word_lower)
        return False, suggestions
    
    def find_suggestions(self, word: str, max_suggestions: int = 3) -> List[str]:
        """Tìm từ gợi ý dựa trên Levenshtein distance"""
        suggestions = []
        
        for dict_word in self.dictionary:
            distance = self.levenshtein_distance(word, dict_word)
            if distance <= 2:  # Ngưỡng khoảng cách
                suggestions.append((dict_word, distance))
        
        # Sắp xếp theo khoảng cách
        suggestions.sort(key=lambda x: x[1])
        return [word for word, _ in suggestions[:max_suggestions]]
    
    @staticmethod
    def levenshtein_distance(s1: str, s2: str) -> int:
        """Tính khoảng cách Levenshtein giữa 2 chuỗi"""
        if len(s1) < len(s2):
            return VietnameseSpellChecker.levenshtein_distance(s2, s1)
        
        if len(s2) == 0:
            return len(s1)
        
        previous_row = range(len(s2) + 1)
        for i, c1 in enumerate(s1):
            current_row = [i + 1]
            for j, c2 in enumerate(s2):
                insertions = previous_row[j + 1] + 1
                deletions = current_row[j] + 1
                substitutions = previous_row[j] + (c1 != c2)
                current_row.append(min(insertions, deletions, substitutions))
            previous_row = current_row
        
        return previous_row[-1]
    
    def check_text(self, text: str) -> Dict:
        """Kiểm tra toàn bộ văn bản"""
        words = re.findall(r'[\wàáạảãâầấậẩẫăằắặẳẵèéẹẻẽêềếệểễìíịỉĩòóọỏõôồốộổỗơờớợởỡùúụủũưừứựửữỳýỵỷỹđ]+', text.lower())
        
        errors = []
        total_words = len(words)
        
        for idx, word in enumerate(words):
            is_correct, suggestions = self.check_word(word)
            if not is_correct:
                errors.append({
                    'word': word,
                    'position': idx,
                    'suggestions': suggestions
                })
        
        return {
            'total_words': total_words,
            'errors': errors,
            'error_rate': len(errors) / total_words if total_words > 0 else 0
        }


class DocumentProcessor:
    """Xử lý tài liệu PDF với OCR và trích xuất nội dung"""
    
    def __init__(self, output_dir='output'):
        # Khởi tạo PaddleOCR
        self.ocr = PaddleOCR(use_angle_cls=True, lang='vi')
        
        # Khởi tạo spell checker
        self.spell_checker = VietnameseSpellChecker()
        
        # Thư mục output
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        os.makedirs(f"{output_dir}/images", exist_ok=True)
        
        # Kết quả
        self.results = {
            'sections': [],
            'images': [],
            'spelling_check': {},
            'metadata': {}
        }
    
    def detect_layout(self, image: np.ndarray) -> List[Dict]:
        """
        Phát hiện layout với bounding boxes
        Sử dụng OCR để phát hiện các vùng văn bản
        """
        layout_elements = []
        
        # OCR để lấy bounding boxes
        ocr_result = self.ocr.ocr(image, cls=True)
        
        if ocr_result is None or len(ocr_result) == 0:
            return layout_elements
        
        for line in ocr_result[0]:
            bbox = line[0]  # [[x1,y1], [x2,y2], [x3,y3], [x4,y4]]
            text = line[1][0]  # Text
            confidence = line[1][1]  # Confidence score
            
            # Chuyển đổi bbox sang format chuẩn
            x_coords = [point[0] for point in bbox]
            y_coords = [point[1] for point in bbox]
            
            x_min, x_max = min(x_coords), max(x_coords)
            y_min, y_max = min(y_coords), max(y_coords)
            
            width = x_max - x_min
            height = y_max - y_min
            
            # Phân loại element type dựa trên kích thước và vị trí
            element_type = self.classify_element(text, width, height, y_min)
            
            layout_elements.append({
                'type': element_type,
                'bbox': {
                    'x': int(x_min),
                    'y': int(y_min),
                    'width': int(width),
                    'height': int(height)
                },
                'text': text,
                'confidence': float(confidence)
            })
        
        return layout_elements
    
    def classify_element(self, text: str, width: float, height: float, y_pos: float) -> str:
        """Phân loại element dựa trên đặc điểm"""
        # Tiêu đề thường to và ở đầu
        if height > 30 and y_pos < 200:
            return 'title'
        
        # Tiêu đề cấp 2 (2.1, 2.2, ...)
        if re.match(r'^\d+\.\d+', text.strip()):
            return 'heading'
        
        # Văn bản thông thường
        return 'text'
    
    def extract_images_from_pdf(self, pdf_path: str) -> List[Dict]:
        """Trích xuất hình ảnh từ PDF với bounding box"""
        doc = fitz.open(pdf_path)
        images = []
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            image_list = page.get_images()
            
            for img_idx, img in enumerate(image_list):
                xref = img[0]
                
                # Lấy thông tin hình ảnh
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                image_ext = base_image["ext"]
                
                # Lưu hình ảnh
                img_filename = f"img_{page_num+1:03d}_{img_idx+1:03d}.{image_ext}"
                img_path = os.path.join(self.output_dir, "images", img_filename)
                
                with open(img_path, "wb") as f:
                    f.write(image_bytes)
                
                # Lấy bounding box của hình ảnh
                img_rects = page.get_image_rects(xref)
                if img_rects:
                    rect = img_rects[0]
                    bbox = {
                        'x': int(rect.x0),
                        'y': int(rect.y0),
                        'width': int(rect.width),
                        'height': int(rect.height)
                    }
                else:
                    bbox = {'x': 0, 'y': 0, 'width': 0, 'height': 0}
                
                images.append({
                    'id': f'img_{page_num+1:03d}_{img_idx+1:03d}',
                    'filename': img_filename,
                    'path': img_path,
                    'page': page_num + 1,
                    'format': image_ext.upper(),
                    'bbox': bbox
                })
        
        doc.close()
        return images
    
    def process_pdf(self, pdf_path: str) -> Dict:
        """Xử lý toàn bộ PDF"""
        print(f"Đang xử lý file: {pdf_path}")
        
        # Trích xuất hình ảnh
        print("Trích xuất hình ảnh...")
        self.results['images'] = self.extract_images_from_pdf(pdf_path)
        print(f"Đã trích xuất {len(self.results['images'])} hình ảnh")
        
        # Chuyển đổi PDF sang hình ảnh để OCR
        doc = fitz.open(pdf_path)
        
        all_text = ""
        section_counter = 1
        
        for page_num in range(len(doc)):
            print(f"Xử lý trang {page_num + 1}/{len(doc)}...")
            
            page = doc[page_num]
            
            # Chuyển page sang image
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # Tăng độ phân giải
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            img_array = np.array(img)
            
            # Detect layout và OCR
            layout_elements = self.detect_layout(img_array)
            
            # Xử lý từng element
            current_section = ""
            
            for element in layout_elements:
                element['page'] = page_num + 1
                
                if element['type'] == 'heading':
                    # Kết thúc section trước nếu có
                    if current_section:
                        self.results['sections'].append({
                            'id': f'section_{section_counter}',
                            'content': current_section.strip() + '\n\n</break>\n',
                            'page': page_num + 1
                        })
                        section_counter += 1
                        current_section = ""
                    
                    current_section = f"\n{element['text']}\n\n"
                else:
                    current_section += element['text'] + " "
                
                all_text += element['text'] + " "
            
            # Lưu section cuối cùng của trang
            if current_section:
                self.results['sections'].append({
                    'id': f'section_{section_counter}',
                    'content': current_section.strip(),
                    'page': page_num + 1
                })
                section_counter += 1
        
        doc.close()
        
        # Kiểm tra chính tả
        print("Kiểm tra chính tả...")
        self.results['spelling_check'] = self.spell_checker.check_text(all_text)
        
        # Metadata
        self.results['metadata'] = {
            'total_pages': len(doc),
            'total_sections': len(self.results['sections']),
            'total_images': len(self.results['images']),
            'total_words': self.results['spelling_check']['total_words'],
            'spelling_errors': len(self.results['spelling_check']['errors'])
        }
        
        return self.results
    
    def export_to_word(self, output_filename='output_document.docx'):
        """Xuất kết quả ra file Word"""
        print(f"Đang xuất file Word: {output_filename}")
        
        doc = Document()
        
        # Thiết lập style
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(13)
        
        # Tiêu đề
        title = doc.add_heading('TÀI LIỆU ĐÃ XỬ LÝ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Thông tin metadata
        doc.add_heading('Thông tin tổng quan', 1)
        metadata_text = f"""
Tổng số trang: {self.results['metadata']['total_pages']}
Tổng số phần: {self.results['metadata']['total_sections']}
Tổng số hình ảnh: {self.results['metadata']['total_images']}
Tổng số từ: {self.results['metadata']['total_words']}
Lỗi chính tả: {self.results['metadata']['spelling_errors']}
        """
        doc.add_paragraph(metadata_text)
        
        # Nội dung các sections
        doc.add_heading('Nội dung', 1)
        
        for section in self.results['sections']:
            # Kiểm tra xem có phải heading không
            content = section['content']
            
            if re.match(r'^\d+\.\d+', content.strip()):
                # Thêm heading
                lines = content.split('\n', 1)
                heading_text = lines[0].strip()
                doc.add_heading(heading_text, 2)
                
                if len(lines) > 1:
                    body_text = lines[1].replace('</break>', '').strip()
                    if body_text:
                        p = doc.add_paragraph(body_text)
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            else:
                # Paragraph thông thường
                text = content.replace('</break>', '').strip()
                if text:
                    p = doc.add_paragraph(text)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Thêm page break nếu có </break>
            if '</break>' in content:
                doc.add_page_break()
        
        # Hình ảnh
        if self.results['images']:
            doc.add_page_break()
            doc.add_heading('Danh sách hình ảnh', 1)
            
            for img in self.results['images']:
                doc.add_heading(f"Hình {img['id']}", 2)
                
                # Thêm hình ảnh vào document
                try:
                    doc.add_picture(img['path'], width=Inches(4))
                except:
                    doc.add_paragraph(f"[Không thể thêm hình ảnh: {img['filename']}]")
                
                # Thông tin hình ảnh
                info_text = f"""
Trang: {img['page']}
Định dạng: {img['format']}
Bounding Box: (x={img['bbox']['x']}, y={img['bbox']['y']}, w={img['bbox']['width']}, h={img['bbox']['height']})
                """
                doc.add_paragraph(info_text)
        
        # Lỗi chính tả
        if self.results['spelling_check']['errors']:
            doc.add_page_break()
            doc.add_heading('Danh sách lỗi chính tả', 1)
            
            for error in self.results['spelling_check']['errors'][:50]:  # Giới hạn 50 lỗi
                p = doc.add_paragraph()
                run = p.add_run(f"Từ: '{error['word']}'")
                run.font.color.rgb = RGBColor(255, 0, 0)
                
                if error['suggestions']:
                    p.add_run(f" → Gợi ý: {', '.join(error['suggestions'])}")
        
        # Lưu file
        output_path = os.path.join(self.output_dir, output_filename)
        doc.save(output_path)
        print(f"Đã lưu file Word: {output_path}")
        
        return output_path


def main():
    """Hàm main để chạy chương trình"""
    
    # Đường dẫn file PDF cần xử lý
    pdf_path = "input_document.pdf"  # Thay đổi đường dẫn này
    
    # Kiểm tra file tồn tại
    if not os.path.exists(pdf_path):
        print(f"Lỗi: Không tìm thấy file {pdf_path}")
        print("Vui lòng đặt file PDF cần xử lý vào thư mục hiện tại")
        return
    
    # Khởi tạo processor
    processor = DocumentProcessor(output_dir='output')
    
    # Xử lý PDF
    results = processor.process_pdf(pdf_path)
    
    # Xuất kết quả
    print("\n" + "="*60)
    print("KẾT QUẢ XỬ LÝ")
    print("="*60)
    print(f"Tổng số trang: {results['metadata']['total_pages']}")
    print(f"Tổng số phần: {results['metadata']['total_sections']}")
    print(f"Tổng số hình ảnh: {results['metadata']['total_images']}")
    print(f"Tổng số từ: {results['metadata']['total_words']}")
    print(f"Lỗi chính tả: {results['metadata']['spelling_errors']}")
    print(f"Tỷ lệ lỗi: {results['spelling_check']['error_rate']:.2%}")
    
    # Xuất ra Word
    output_file = processor.export_to_word('tai_lieu_da_xu_ly.docx')
    
    print("\n" + "="*60)
    print(f"Hoàn thành! File Word đã được lưu tại: {output_file}")
    print("="*60)


if __name__ == "__main__":
    main()