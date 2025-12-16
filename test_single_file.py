"""
Test xử lý 1 file PDF đơn (đã cải thiện xử lý tiếng Việt sau OCR)
"""

import os
import re
import unicodedata
from typing import Dict, List, Any, Tuple, Optional

from document_processor import EnhancedDocumentProcessor

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


# =========================
# 1) VIETNAMESE POSTPROCESS
# =========================

def normalize_nfc(text: str) -> str:
    """Chuẩn hoá Unicode về NFC để tránh rơi dấu/combining marks."""
    if not text:
        return text
    return unicodedata.normalize("NFC", text)


def vi_ocr_cleanup(text: str) -> str:
    """
    Làm sạch lỗi OCR thường gặp (nhẹ nhàng, hạn chế phá chữ).
    - normalize NFC
    - chuẩn hoá khoảng trắng
    - sửa dính dấu câu
    """
    if not text:
        return text

    text = normalize_nfc(text)

    # chuẩn hoá xuống dòng
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # khoảng trắng thừa
    text = re.sub(r"[ \t]+", " ", text)

    # dấu câu dính vào từ
    text = re.sub(r"\s+([,.;:!?])", r"\1", text)          # "từ ," -> "từ,"
    text = re.sub(r"([,.;:!?])([^\s])", r"\1 \2", text)   # ",từ" -> ", từ"

    # nhiều dòng trống
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


def build_full_text_from_sections(sections: List[Dict[str, Any]]) -> str:
    """Ghép text từ sections và cleanup."""
    parts = []
    for sec in sections:
        content = sec.get("content", "")
        content = vi_ocr_cleanup(content)
        if content:
            parts.append(content)
    return "\n\n".join(parts).strip()


# =====================================
# 2) SPELL CORRECTION: POSITION FIRST
# =====================================

def apply_corrections_by_position(full_text: str, errors: List[Dict[str, Any]]) -> Tuple[str, int]:
    """
    Sửa theo vị trí (position) an toàn hơn regex replace.
    Yêu cầu: error['position'] là index trong full_text, và error['word'] khớp substring.
    """
    if not full_text or not errors:
        return full_text, 0

    # sort giảm dần để không lệch index
    errors_sorted = sorted(
        [e for e in errors if isinstance(e, dict) and e.get("position") is not None],
        key=lambda e: e.get("position", -1),
        reverse=True
    )

    text = full_text
    fixed = 0

    for e in errors_sorted:
        pos = e.get("position", None)
        wrong = e.get("word", "")
        suggs = e.get("suggestions", []) or []
        if pos is None or pos < 0 or not wrong or not suggs:
            continue

        correct = suggs[0]
        wrong = normalize_nfc(str(wrong))
        correct = normalize_nfc(str(correct))

        # kiểm tra substring tại vị trí có đúng "wrong" không
        segment = text[pos:pos + len(wrong)]
        if normalize_nfc(segment) != wrong:
            # nếu không khớp, bỏ qua để tránh thay nhầm
            continue

        text = text[:pos] + correct + text[pos + len(wrong):]
        fixed += 1

    return text, fixed


def apply_corrections_by_regex(full_text: str, errors: List[Dict[str, Any]]) -> Tuple[str, int]:
    """
    Fallback: sửa bằng regex word boundary.
    Với tiếng Việt: dùng boundary kiểu "không phải chữ/số/_" để đỡ sai.
    """
    if not full_text or not errors:
        return full_text, 0

    text = full_text
    fixed = 0

    for e in errors:
        wrong = normalize_nfc(str(e.get("word", "") or ""))
        suggs = e.get("suggestions", []) or []
        if not wrong or not suggs:
            continue

        correct = normalize_nfc(str(suggs[0]))

        # boundary cho Unicode (tránh \b hơi lạ với dấu)
        pattern = rf"(?<![\wÀ-ỹ]){re.escape(wrong)}(?![\wÀ-ỹ])"
        new_text, n = re.subn(pattern, correct, text, flags=re.IGNORECASE)
        if n > 0:
            fixed += n
            text = new_text

    return text, fixed


def auto_correct_spelling(results: Dict[str, Any]) -> Tuple[str, int]:
    """
    Tự động sửa lỗi:
    - Ghép + cleanup
    - Ưu tiên sửa theo position (nếu position map đúng full_text)
    - Fallback regex
    """
    full_text = build_full_text_from_sections(results.get("sections", []))
    errors = results.get("spelling_check", {}).get("errors", []) or []

    # thử sửa theo position trước
    corrected, fixed_pos = apply_corrections_by_position(full_text, errors)

    # nếu không sửa được gì (hoặc quá ít), fallback regex
    if fixed_pos == 0 and errors:
        corrected, fixed_rx = apply_corrections_by_regex(full_text, errors)
        return corrected, fixed_rx

    return corrected, fixed_pos


# =========================
# 3) WORD EXPORT VI SAFE
# =========================

def set_word_font_vi_safe(doc: Document, font_name: str = "Times New Roman", font_size_pt: int = 13) -> None:
    """
    Set font đầy đủ để Word không fallback gây lỗi ký tự tiếng Việt.
    """
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(font_size_pt)

    rFonts = style._element.rPr.rFonts
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:cs"), font_name)


def export_corrected_word(corrected_text: str, output_dir: str, filename: str, results: Dict[str, Any], fixed_count: int) -> str:
    """Xuất văn bản đã sửa ra Word (an toàn tiếng Việt)."""
    doc = Document()
    set_word_font_vi_safe(doc, font_name="Times New Roman", font_size_pt=13)

    # Tiêu đề
    title = doc.add_heading("TÀI LIỆU ĐÃ HẬU XỬ LÝ (TIẾNG VIỆT)", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Thông tin
    doc.add_heading("Thông tin", 1)
    info_text = (
        f"Số lỗi đã sửa (ước tính): {fixed_count}\n"
        f"Tổng số từ: {results.get('metadata', {}).get('total_words', 'N/A')}\n"
        f"Tỷ lệ lỗi ban đầu: {results.get('spelling_check', {}).get('error_rate', 0):.2%}\n"
        f"Gợi ý: Nếu còn rụng chữ nặng (vd: 'th', 'phn'), hãy chỉnh OCR lang/model trong document_processor.py.\n"
    )
    doc.add_paragraph(info_text)

    # Nội dung
    doc.add_heading("Nội dung", 1)

    # Chia đoạn theo double newline
    paragraphs = corrected_text.split("\n\n")
    for para in paragraphs:
        para = para.strip()
        if para:
            p = doc.add_paragraph(para)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    output_path = os.path.join(output_dir, filename)
    doc.save(output_path)
    return output_path


# =========================
# 4) MAIN
# =========================

def main():
    print("=" * 80)
    print("🧪 TEST XỬ LÝ 1 FILE PDF (có hậu xử lý tiếng Việt)")
    print("=" * 80)

    pdf_file = "input_pdfs/ban_thao_van_dap.pdf"  # File nhỏ nhất - 83 trang

    if not os.path.exists(pdf_file):
        print(f"❌ Không tìm thấy file: {pdf_file}")
        print("\n📝 Vui lòng:")
        print("1. Kiểm tra tên file đúng chưa")
        print("2. Kiểm tra file có trong thư mục input_pdfs/ không")
        return

    print(f"\n📄 File test: {os.path.basename(pdf_file)}")

    output_dir = "output_test_single"
    os.makedirs(output_dir, exist_ok=True)
    print(f"📁 Kết quả sẽ lưu trong: {output_dir}/")

    confirm = input("\n▶️  Bắt đầu xử lý? (y/n): ")
    if confirm.lower() != "y":
        print("❌ Đã hủy")
        return

    try:
        processor = EnhancedDocumentProcessor(output_dir=output_dir)

        print("\n⏳ Đang xử lý...")
        results = processor.process_pdf(pdf_file, debug=False)

        # Normalize + cleanup sections ngay sau OCR (quan trọng)
        for sec in results.get("sections", []):
            sec["content"] = vi_ocr_cleanup(sec.get("content", ""))

        output_filename = f"{os.path.basename(pdf_file).replace('.pdf', '')}_processed.docx"
        output_path = processor.export_to_word(output_filename)

        print("\n" + "=" * 80)
        print("🔧 HẬU XỬ LÝ + SỬA (NẾU CÓ) LỖI CHÍNH TẢ")
        print("=" * 80)

        errors = results.get("spelling_check", {}).get("errors", []) or []
        if errors:
            print(f"\n📝 Tìm thấy {len(errors)} lỗi (theo spell checker hiện tại)")
            fix_spelling = input("\n▶️  Bạn có muốn tự động sửa + làm sạch tiếng Việt? (y/n): ")

            if fix_spelling.lower() == "y":
                corrected_text, fixed_count = auto_correct_spelling(results)

                corrected_file = os.path.join(output_dir, "ocr_result_corrected.txt")
                with open(corrected_file, "w", encoding="utf-8") as f:
                    f.write("=" * 80 + "\n")
                    f.write(f"KẾT QUẢ OCR ĐÃ HẬU XỬ LÝ: {os.path.basename(pdf_file)}\n")
                    f.write("=" * 80 + "\n\n")
                    f.write(corrected_text)

                print(f"\n✅ Đã sửa (ước tính): {fixed_count} thay thế")
                print(f"📄 File đã sửa: {corrected_file}")

                corrected_word_file = f"{os.path.basename(pdf_file).replace('.pdf', '')}_corrected.docx"
                corrected_word_path = export_corrected_word(
                    corrected_text, output_dir, corrected_word_file, results, fixed_count
                )
                print(f"💾 File Word đã sửa: {corrected_word_path}")
            else:
                print("⏭️  Bỏ qua hậu xử lý")
        else:
            print("\n✅ Spell checker báo: Không có lỗi chính tả!")

        print("\n" + "=" * 80)
        print("✅ HOÀN THÀNH!")
        print("=" * 80)

        md = results.get("metadata", {})
        print(f"\n📊 Thống kê:")
        print(f"   📄 Tổng số trang: {md.get('total_pages', 'N/A')}")
        print(f"   📝 Tổng số phần: {md.get('total_sections', 'N/A')}")
        print(f"   🖼️  Tổng số hình: {md.get('total_images', 'N/A')}")
        print(f"   📖 Tổng số từ: {md.get('total_words', 'N/A')}")
        print(f"   ❌ Lỗi chính tả: {md.get('spelling_errors', 'N/A')}")

        print(f"\n📝 PREVIEW NỘI DUNG OCR (5 phần đầu):")
        print("-" * 80)
        for i, section in enumerate(results.get("sections", [])[:5], 1):
            content_preview = (section.get("content", "")[:200]).replace("\n", " ")
            print(f"\n{i}. Trang {section.get('page', 'N/A')}:")
            print(f"   {content_preview}...")

        print(f"\n🖼️  DANH SÁCH HÌNH ẢNH ({len(results.get('images', []))} ảnh):")
        print("-" * 80)
        for i, img in enumerate(results.get("images", [])[:10], 1):
            print(f"{i}. {img.get('filename')}")
            print(
                f"   Trang: {img.get('page')} | Format: {img.get('format')} | Path: {img.get('path')}"
            )

        if len(results.get("images", [])) > 10:
            print(f"   ... và {len(results.get('images', [])) - 10} hình ảnh khác")

        # Lưu OCR text
        ocr_text_file = os.path.join(output_dir, "ocr_result.txt")
        with open(ocr_text_file, "w", encoding="utf-8") as f:
            f.write("=" * 80 + "\n")
            f.write(f"KẾT QUẢ OCR: {os.path.basename(pdf_file)}\n")
            f.write("=" * 80 + "\n\n")

            for section in results.get("sections", []):
                f.write(f"\n{'='*60}\n")
                f.write(f"Trang {section.get('page')} - Section {section.get('id')}\n")
                f.write(f"{'='*60}\n")
                f.write(section.get("content", ""))
                f.write("\n\n")

        print(f"\n📄 File text OCR: {ocr_text_file}")

        # Lưu danh sách hình
        images_list_file = os.path.join(output_dir, "images_list.txt")
        with open(images_list_file, "w", encoding="utf-8") as f:
            f.write("=" * 80 + "\n")
            f.write(f"DANH SÁCH HÌNH ẢNH: {os.path.basename(pdf_file)}\n")
            f.write(f"Tổng số: {len(results.get('images', []))} hình ảnh\n")
            f.write("=" * 80 + "\n\n")

            for i, img in enumerate(results.get("images", []), 1):
                f.write(f"{i}. {img.get('filename')}\n")
                f.write(f"   ID: {img.get('id')}\n")
                f.write(f"   Trang: {img.get('page')}\n")
                f.write(f"   Format: {img.get('format')}\n")
                f.write(f"   Path: {img.get('path')}\n")
                bbox = img.get("bbox", {}) or {}
                f.write(
                    f"   Bbox: x={bbox.get('x')}, y={bbox.get('y')}, "
                    f"w={bbox.get('width')}, h={bbox.get('height')}\n"
                )
                f.write("\n")

        print(f"🖼️  File danh sách ảnh: {images_list_file}")

        # Stats
        stats_file = os.path.join(output_dir, "statistics.txt")
        with open(stats_file, "w", encoding="utf-8") as f:
            f.write("=" * 80 + "\n")
            f.write(f"THỐNG KÊ CHI TIẾT: {os.path.basename(pdf_file)}\n")
            f.write("=" * 80 + "\n\n")

            f.write("TỔNG QUAN:\n")
            f.write(f"  Tổng số trang: {md.get('total_pages', 'N/A')}\n")
            f.write(f"  Tổng số phần: {md.get('total_sections', 'N/A')}\n")
            f.write(f"  Tổng số hình ảnh: {md.get('total_images', 'N/A')}\n")
            f.write(f"  Tổng số từ: {md.get('total_words', 'N/A')}\n")
            f.write(f"  Lỗi chính tả: {md.get('spelling_errors', 'N/A')}\n")
            f.write(f"  Tỷ lệ lỗi: {results.get('spelling_check', {}).get('error_rate', 0):.2%}\n\n")

            if errors:
                f.write("LỖI CHÍNH TẢ (20 lỗi đầu):\n")
                for i, error in enumerate(errors[:20], 1):
                    f.write(f"  {i}. '{error.get('word')}' (vị trí: {error.get('position')})\n")
                    if error.get("suggestions"):
                        f.write(f"     Gợi ý: {', '.join(error.get('suggestions'))}\n")

        print(f"📊 File thống kê: {stats_file}")

        print(f"\n💾 File Word (raw export): {output_path}")
        print(f"📁 Thư mục hình ảnh: {output_dir}/images/")
        print("=" * 80)

    except Exception as e:
        print(f"\n❌ LỖI: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()
